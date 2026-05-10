#!/usr/bin/env python3
"""
RemarkableAnnotate — extract highlighted text from reMarkable PDFs via USB web interface.

The reMarkable USB web interface runs at http://10.11.99.1 when the device is connected
via USB. No SSH, no Developer Mode, no password required.

Usage:
  python3 extract.py list    <host>
  python3 extract.py extract <host> <uuid> <output_path>
"""

import sys
import json
import io
import re
import zipfile
import traceback


# ---------------------------------------------------------------------------
# HTTP helpers
# ---------------------------------------------------------------------------

def get_json(host, path):
    import requests
    resp = requests.get(f"http://{host}{path}", timeout=10)
    resp.raise_for_status()
    return resp.json()


def download_bytes(host, path, timeout=120):
    import requests
    resp = requests.get(f"http://{host}{path}", timeout=timeout, stream=True)
    resp.raise_for_status()
    buf = io.BytesIO()
    for chunk in resp.iter_content(chunk_size=65536):
        buf.write(chunk)
    buf.seek(0)
    return buf


# ---------------------------------------------------------------------------
# .rmdoc zip helpers
# ---------------------------------------------------------------------------

def read_zip_json(zf, name):
    try:
        with zf.open(name) as f:
            return json.load(f)
    except Exception:
        return None


def _parse_idx(v, fallback):
    """Parse a reMarkable page index — stored as a hex string in firmware."""
    if isinstance(v, int):
        return v
    try:
        return int(str(v), 16)
    except (ValueError, TypeError):
        try:
            return int(str(v))
        except (ValueError, TypeError):
            return fallback


def page_list(content):
    """Return list of {id, pdf_index} from .content JSON."""
    c_pages = content.get("cPages", {}).get("pages", [])
    if c_pages:
        result = []
        for i, p in enumerate(c_pages):
            raw_idx = p.get("idx", {}).get("value", i) if isinstance(p.get("idx"), dict) else i
            result.append({
                "id": p.get("id", ""),
                "pdf_index": _parse_idx(raw_idx, i),
            })
        return result
    old = content.get("pages", [])
    return [{"id": pid, "pdf_index": i} for i, pid in enumerate(old)]


# ---------------------------------------------------------------------------
# PDF page lookup
# ---------------------------------------------------------------------------

def _find_physical_page(pdf_doc, pdf_index_hint, glyph_texts):
    """Return (fitz.Page, physical_idx) for the page that contains the glyph text.

    The .content idx value is the printed page number, NOT the 0-based physical
    page in the PDF file.  We use PyMuPDF's page-label lookup first, then a
    linear scan so we never walk 150+ pages outward from a wrong hint.
    """
    if pdf_doc is None:
        return None, pdf_index_hint
    n = pdf_doc.page_count

    if not glyph_texts:
        idx = max(0, min(pdf_index_hint, n - 1))
        return pdf_doc[idx], idx

    # Try page-label lookup: the printed page number may be stored as a PDF label
    try:
        label_pages = pdf_doc.get_page_numbers(str(pdf_index_hint + 1))
        if label_pages:
            return pdf_doc[label_pages[0]], label_pages[0]
    except Exception:
        pass

    # Build a search phrase from the first glyph text — avoid words with apostrophes
    words = glyph_texts[0].strip().split()
    safe = [w for w in words if "'" not in w and "’" not in w][:8]
    if len(safe) < 3:
        safe = words[:6]
    search = " ".join(safe[:6])

    # Linear scan (finds page 82 after ~83 checks, faster than expanding from 238)
    for idx in range(n):
        if pdf_doc[idx].search_for(search):
            return pdf_doc[idx], idx

    # Give up — return the hint page
    idx = max(0, min(pdf_index_hint, n - 1))
    return pdf_doc[idx], idx


# ---------------------------------------------------------------------------
# .rm highlight parsing (rmscene)
# ---------------------------------------------------------------------------

def _is_highlighter_pen(line):
    # rmscene Line uses 'tool'; older versions used 'pen'
    pen = getattr(line, 'tool', None) or getattr(line, 'pen', None)
    if pen is None:
        return False
    try:
        from rmscene.scene_items import Pen
        hi_vals = set()
        for name in ('HIGHLIGHTER_1', 'HIGHLIGHTER_2', 'HIGHLIGHTER'):
            p = getattr(Pen, name, None)
            if p is not None:
                hi_vals.add(int(p))
        if not hi_vals:
            hi_vals = {5, 18}
        return int(pen) in hi_vals
    except Exception:
        return int(pen) in (5, 18)


def _calibrate_rm_to_pdf(blocks, pdf_page):
    """Derive the rm-space → pdf-space affine transform from GlyphRange rectangles."""
    try:
        from rmscene import SceneGlyphItemBlock
        from rmscene.scene_items import PenColor

        rm_xs, rm_ys, pdf_xs, pdf_ys = [], [], [], []

        for block in blocks:
            if not isinstance(block, SceneGlyphItemBlock):
                continue
            if block.item is None or block.item.value is None:
                continue
            glyph = block.item.value
            if glyph.color != PenColor.HIGHLIGHT or not glyph.text:
                continue

            rects = getattr(glyph, 'rectangles', None) or []
            if not rects:
                continue

            words = glyph.text.strip().split()
            if len(words) < 2:
                continue

            found = None
            for n in (5, 4, 3, 2):
                if len(words) >= n:
                    hits = pdf_page.search_for(" ".join(words[:n]))
                    if hits:
                        found = hits[0]
                        break
            if found is None:
                continue

            r = rects[0]
            if hasattr(r, 'x') and hasattr(r, 'y'):
                rx, ry, rw, rh = r.x, r.y, r.width, r.height
            elif hasattr(r, 'pos'):
                rx, ry = r.pos.x, r.pos.y
                rw, rh = r.width, r.height
            else:
                continue

            rm_xs.extend([rx, rx + rw])
            rm_ys.extend([ry, ry + rh])
            pdf_xs.extend([found.x0, found.x1])
            pdf_ys.extend([found.y0, found.y1])

        if len(rm_xs) < 4:
            return None

        def fit(rm_vals, pdf_vals):
            n = len(rm_vals)
            rm_mean = sum(rm_vals) / n
            pdf_mean = sum(pdf_vals) / n
            num = sum((rm_vals[i] - rm_mean) * (pdf_vals[i] - pdf_mean) for i in range(n))
            den = sum((rm_vals[i] - rm_mean) ** 2 for i in range(n))
            if den < 1e-10:
                return None
            s = num / den
            return s, pdf_mean - s * rm_mean

        xfit = fit(rm_xs, pdf_xs)
        yfit = fit(rm_ys, pdf_ys)
        if xfit is None or yfit is None:
            return None

        return (xfit[0], xfit[1], yfit[0], yfit[1])

    except Exception:
        return None


def _pdf_text_in_rm_rect(pdf_page, rm_rect, transform=None):
    """Extract text from a PDF area mapped from a reMarkable stroke bounding box."""
    try:
        import fitz
        x1, y1, x2, y2 = rm_rect
        pr = pdf_page.rect

        if transform is not None:
            xs, xo, ys, yo = transform
            pdf_x1 = x1 * xs + xo
            pdf_y1 = y1 * ys + yo
            pdf_x2 = x2 * xs + xo
            pdf_y2 = y2 * ys + yo
        else:
            RM_W, RM_H = 1404.0, 1872.0
            scale = min(RM_W / pr.width, RM_H / pr.height)
            xm = (RM_W - pr.width * scale) / 2
            ym = (RM_H - pr.height * scale) / 2
            pdf_x1 = (x1 - xm) / scale
            pdf_y1 = (y1 - ym) / scale
            pdf_x2 = (x2 - xm) / scale
            pdf_y2 = (y2 - ym) / scale

        pad = 12.0
        clip = fitz.Rect(
            max(0.0, pdf_x1 - 5),
            max(0.0, pdf_y1 - pad),
            min(pr.width, pdf_x2 + 5),
            min(pr.height, pdf_y2 + pad),
        )
        text = pdf_page.get_text("text", clip=clip).strip()
        return " ".join(text.split()) if text else ""
    except Exception:
        return ""


def _norm_word(w):
    """Lowercase and strip all non-alphanumeric chars for fuzzy matching."""
    return "".join(c for c in w.lower() if c.isalnum())


def _glyph_end_pdf_y(pdf_page, glyph_texts):
    """Return the bottom-Y (PDF points) of the last glyph text on the page, or None."""
    if not glyph_texts:
        return None
    last_words = glyph_texts[-1].strip().split()
    safe = [w for w in last_words if "'" not in w and "’" not in w and w.isalpha()]
    for n in range(min(5, len(safe)), 1, -1):
        hits = pdf_page.search_for(" ".join(safe[-n:]))
        if hits:
            return max(h.y1 for h in hits)
    return None


def _text_before_glyph(pdf_page, glyph_texts, stroke_rect):
    """Find text immediately before the first glyph word in the PDF word stream.

    Used when a stroke sits above/before the glyph region — typically a word
    or two at the very start of a highlight that the GlyphRange missed.
    Words joined to the found token by em-dashes have the prefix stripped
    (e.g. "shortly—VCs" → "VCs").
    """
    try:
        if not glyph_texts:
            return ""
        first_words = glyph_texts[0].strip().split()
        safe = [_norm_word(w) for w in first_words[:4] if _norm_word(w)]
        if len(safe) < 2:
            return ""

        pdf_words = pdf_page.get_text("words")
        pdf_norms = [_norm_word(w[4]) for w in pdf_words]

        # Find first occurrence of glyph start in the word stream
        glyph_start = -1
        for i in range(len(pdf_norms) - len(safe) + 1):
            if pdf_norms[i:i + len(safe)] == safe:
                glyph_start = i
                break
        if glyph_start <= 0:
            return ""

        # How many words to look back: estimate from stroke width
        rx1, ry1, rx2, ry2 = stroke_rect
        look_back = max(1, int(max(rx2 - rx1, 1) / 100) + 1)

        result = []
        for w_tuple in pdf_words[max(0, glyph_start - look_back):glyph_start]:
            word = w_tuple[4]
            # Strip any prefix joined by em-dash (e.g. "shortly—VCs" → "VCs")
            parts = word.split("—")
            clean = parts[-1].strip()
            if clean:
                result.append(clean)
        return " ".join(result)
    except Exception:
        return ""


def _stroke_is_after_glyph(pdf_page, glyph_texts, stroke_rect):
    """True when the stroke's centre Y is plausibly below the glyph text region.

    Uses a proportional rm→PDF Y estimate (good enough for above/below checks)
    to avoid calling _text_after_glyph for strokes that overlap the glyph text.
    """
    end_y = _glyph_end_pdf_y(pdf_page, glyph_texts)
    if end_y is None:
        return True  # no reference point — allow fallback
    _, ry1, _, ry2 = stroke_rect
    pr = pdf_page.rect
    stroke_pdf_y = ((ry1 + ry2) / 2) * pr.height / 1872.0
    return stroke_pdf_y > end_y - 20  # 20pt tolerance


def _text_after_glyph(pdf_page, glyph_texts, stroke_rect=None):
    """Extract text immediately following the last glyph highlight.

    Uses word-by-word normalised matching so curly apostrophes and other
    encoding differences between the rmscene text and the PDF don't matter.
    stroke_rect (x1,y1,x2,y2) in rm-space is used to estimate max word count.
    """
    try:
        if not glyph_texts:
            return ""
        last_words = glyph_texts[-1].strip().split()
        if len(last_words) < 3:
            return ""

        # Estimate how many words the stroke covers from its rm-space width.
        # rm canvas is 1404 units wide; assume ~100 units per word.
        if stroke_rect is not None:
            rx1, ry1, rx2, ry2 = stroke_rect
            rw = max(rx2 - rx1, 1)
            max_words = max(3, int(rw / 100) + 2)
        else:
            max_words = 15

        # Get all words from the PDF page with their positions
        pdf_words = pdf_page.get_text("words")  # (x0,y0,x1,y1,word,...)
        if not pdf_words:
            return ""

        pdf_norms = [_norm_word(w[4]) for w in pdf_words]
        target_tail = [_norm_word(w) for w in last_words[-5:]]
        target_tail = [t for t in target_tail if t]  # drop empty after stripping

        if not target_tail:
            return ""

        # Find the last occurrence of the tail sequence in the PDF word list
        best_end = -1
        tlen = len(target_tail)
        for i in range(len(pdf_norms) - tlen, -1, -1):
            if pdf_norms[i:i + tlen] == target_tail:
                best_end = i + tlen - 1
                break

        if best_end < 0:
            return ""

        # Collect words after the match, up to max_words or sentence end
        result = []
        for w_tuple in pdf_words[best_end + 1: best_end + 1 + max_words]:
            word = w_tuple[4]
            result.append(word)
            if word.endswith(('.', '!', '?')):
                break
        return " ".join(result)
    except Exception:
        return ""


def highlight_texts(rm_bytes, pdf_doc=None, pdf_index_hint=0):
    """Return (texts, physical_idx) from a .rm annotation file.

    physical_idx is the 0-based physical PDF page index (not the printed page
    number stored in .content idx).  Combines two methods:
    - SceneGlyphItemBlock (GlyphRange): direct text, no PDF needed.
    - SceneLineItemBlock with HIGHLIGHTER pen: uses the known glyph text to
      locate the correct physical PDF page, then extracts text using calibrated
      coordinate mapping or a text-search fallback.
    """
    try:
        from rmscene import read_blocks, SceneGlyphItemBlock, SceneLineItemBlock
        from rmscene.scene_items import PenColor
        blocks = list(read_blocks(io.BytesIO(rm_bytes)))
    except Exception:
        return [], pdf_index_hint

    # --- GlyphRange highlights (direct text) ---
    ranges = []
    for block in blocks:
        if not isinstance(block, SceneGlyphItemBlock):
            continue
        if block.item is None or block.item.value is None:
            continue
        glyph = block.item.value
        if glyph.color == PenColor.HIGHLIGHT and glyph.text:
            ranges.append((glyph.start, glyph.length, glyph.text.strip()))

    glyph_texts = []
    if ranges:
        ranges.sort(key=lambda r: r[0])
        cur_start, cur_len, cur_text = ranges[0]
        for start, length, text in ranges[1:]:
            gap = start - (cur_start + cur_len)
            if gap <= 5:
                if gap < 0:
                    # Overlapping ranges — trim the already-covered chars from
                    # the new fragment to avoid duplicated words ("in the in the")
                    text = text[abs(gap):]
                cur_text = (cur_text + " " + text).strip()
                cur_len = (start + length) - cur_start
            else:
                glyph_texts.append(cur_text)
                cur_start, cur_len, cur_text = start, length, text
        glyph_texts.append(cur_text)

    # --- Determine physical page index ---
    physical_idx = pdf_index_hint
    pdf_page = None
    if pdf_doc is not None and glyph_texts:
        pdf_page, physical_idx = _find_physical_page(pdf_doc, pdf_index_hint, glyph_texts)

    # --- Stroke-based highlights ---
    stroke_texts = []
    if pdf_doc is not None:
        # Detect HIGHLIGHTER strokes on this page
        stroke_rects = []
        for block in blocks:
            if not isinstance(block, SceneLineItemBlock):
                continue
            if block.item is None or block.item.value is None:
                continue
            line = block.item.value
            if _is_highlighter_pen(line):
                pts = getattr(line, 'points', None) or []
                if pts:
                    xs = [p.x for p in pts]
                    ys = [p.y for p in pts]
                    stroke_rects.append((min(xs), min(ys), max(xs), max(ys)))

        if stroke_rects:
            # Use already-discovered physical page, or find it now if no glyphs
            if pdf_page is None:
                pdf_page, physical_idx = _find_physical_page(pdf_doc, pdf_index_hint, glyph_texts)

            if pdf_page is not None:
                transform = _calibrate_rm_to_pdf(blocks, pdf_page)
                stroke_rects.sort(key=lambda r: r[1])

                seen_after = set()
                seen_before = set()
                for rect in stroke_rects:
                    t = ""
                    if transform is not None:
                        t = _pdf_text_in_rm_rect(pdf_page, rect, transform)
                    if not t:
                        if _stroke_is_after_glyph(pdf_page, glyph_texts, rect):
                            t = _text_after_glyph(pdf_page, glyph_texts, stroke_rect=rect)
                            if t and t not in seen_after:
                                seen_after.add(t)
                                stroke_texts.append(t)
                        else:
                            # Stroke is before/at the glyph start — may be words
                            # the GlyphRange missed at the beginning of the selection
                            prefix = _text_before_glyph(pdf_page, glyph_texts, rect)
                            if prefix and prefix not in seen_before and glyph_texts:
                                seen_before.add(prefix)
                                glyph_texts[0] = prefix + " " + glyph_texts[0]

    return glyph_texts + stroke_texts, physical_idx


# ---------------------------------------------------------------------------
# list command
# ---------------------------------------------------------------------------

def cmd_list(host):
    try:
        tree = _build_tree(host, "")
    except Exception as e:
        msg = str(e).lower()
        if any(kw in msg for kw in ("connection", "timeout", "refused", "unreachable", "network")):
            _err("Could not connect to reMarkable. Wake it up and make sure it is connected via USB.")
        else:
            _err(f"Failed to list documents: {e}")
        return
    _ok({"tree": tree})


def _build_tree(host, folder_id):
    """Return a list of nodes: folders have children[], documents have uuid."""
    path = f"/documents/{folder_id}" if folder_id else "/documents/"
    try:
        items = get_json(host, path)
    except Exception:
        if not folder_id:
            raise  # Root fetch failed — real connection error, not just a missing subfolder
        return []

    def sort_key(i):
        return (i.get("VissibleName") or "").lower()

    nodes = []
    for item in sorted(items, key=sort_key):
        t = (item.get("Type") or "").lower()
        title = item.get("VissibleName") or "Untitled"
        uuid = item.get("ID") or ""
        if not uuid:
            continue
        if "collection" in t:
            nodes.append({
                "kind": "folder",
                "uuid": uuid,
                "title": title,
                "children": _build_tree(host, uuid),
            })
        else:
            nodes.append({
                "kind": "document",
                "uuid": uuid,
                "title": title,
            })
    return nodes


# ---------------------------------------------------------------------------
# extract command
# ---------------------------------------------------------------------------

def cmd_extract(host, uuid, output_path, title="Untitled"):
    rmdoc_buf = download_bytes(host, f"/download/{uuid}/rmdoc")

    with zipfile.ZipFile(rmdoc_buf) as zf:
        names = zf.namelist()

        content_name = _find(names, lambda n: n.endswith(".content"))
        if not content_name:
            _err("No .content file found in rmdoc bundle")

        content = read_zip_json(zf, content_name)
        if not content:
            _err("Could not read .content file")

        prefix = content_name[:-len(".content")]
        pages = page_list(content)

        # Open embedded PDF for stroke-based highlight extraction (optional)
        pdf_doc = None
        pdf_name = _find(names, lambda n: n.endswith(".pdf"))
        if pdf_name:
            try:
                import fitz
                pdf_doc = fitz.open(stream=zf.read(pdf_name), filetype="pdf")
            except Exception:
                pass

        # If PDF not embedded, try downloading it separately
        if pdf_doc is None:
            try:
                import fitz
                pdf_buf = download_bytes(host, f"/download/{uuid}/pdf")
                pdf_doc = fitz.open(stream=pdf_buf.read(), filetype="pdf")
            except Exception:
                pass

        hits = []
        for pg in pages:
            rm_name = f"{prefix}/{pg['id']}.rm"
            if rm_name not in names:
                rm_name = f"{pg['id']}.rm"
            if rm_name not in names:
                continue

            with zf.open(rm_name) as f:
                rm_data = f.read()

            texts, physical_idx = highlight_texts(rm_data, pdf_doc=pdf_doc,
                                                  pdf_index_hint=pg["pdf_index"])
            pdf_page_num = physical_idx + 1

            pdf_words = []
            if pdf_doc is not None and texts:
                try:
                    pdf_words = pdf_doc[physical_idx].get_text("words")
                except Exception:
                    pass

            for text in texts:
                ctx_before, ctx_after = _context_around(pdf_words, text)
                hits.append((f"Page {pdf_page_num}", text, ctx_before, ctx_after))

        if pdf_doc is not None:
            try:
                pdf_doc.close()
            except Exception:
                pass

    _write_docx(output_path, title, hits)
    _ok({"highlight_count": len(hits)})


def _find(names, predicate):
    return next((n for n in names if predicate(n)), None)


def _context_around(pdf_words, highlight_text, n_sentences=2, max_words=50):
    """Return (before, after) context strings surrounding a highlight.

    Searches the PDF word list for the start and end of the highlight text,
    then walks outward to collect up to n_sentences (capped at max_words).
    """
    try:
        if not pdf_words or not highlight_text:
            return "", ""

        pdf_norms = [_norm_word(w[4]) for w in pdf_words]
        hl_words = highlight_text.strip().split()

        start_safe = [_norm_word(w) for w in hl_words[:5] if _norm_word(w)]
        end_safe   = [_norm_word(w) for w in hl_words[-5:] if _norm_word(w)]
        if len(start_safe) < 2 or len(end_safe) < 2:
            return "", ""

        hl_start = next(
            (i for i in range(len(pdf_norms) - len(start_safe) + 1)
             if pdf_norms[i:i + len(start_safe)] == start_safe), -1)
        hl_end = next(
            (i + len(end_safe) - 1
             for i in range(len(pdf_norms) - len(end_safe), -1, -1)
             if pdf_norms[i:i + len(end_safe)] == end_safe), -1)

        if hl_start < 0 or hl_end < 0:
            return "", ""

        # --- before context ---
        window = pdf_words[max(0, hl_start - max_words):hl_start]
        count = 0
        trim = 0
        for i in range(len(window) - 1, -1, -1):
            if window[i][4].rstrip('"\'').endswith(('.', '!', '?')):
                count += 1
                if count >= n_sentences:
                    trim = i + 1
                    break
        before = "…" + " ".join(w[4] for w in window[trim:]) if window[trim:] else ""

        # --- after context ---
        window = pdf_words[hl_end + 1:hl_end + 1 + max_words]
        result, count = [], 0
        for w_tuple in window:
            word = w_tuple[4]
            result.append(word)
            if word.rstrip('"\'').endswith(('.', '!', '?')):
                count += 1
                if count >= n_sentences:
                    break
        after = (" ".join(result) + "…") if result else ""

        return before, after
    except Exception:
        return "", ""


def _write_docx(path, title, hits):
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def add_left_border(paragraph, color="4472C4"):
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "18")
        left.set(qn("w:space"), "4")
        left.set(qn("w:color"), color)
        pBdr.append(left)
        pPr.append(pBdr)

    doc = Document()

    # Tighten default paragraph spacing
    doc.styles["Normal"].paragraph_format.space_after = Pt(4)

    doc.add_heading(title, 0)

    current_label = None
    for label, text, ctx_before, ctx_after in hits:
        if label is not None and label != current_label:
            current_label = label
            doc.add_heading(label, 2)

        if ctx_before:
            p = doc.add_paragraph()
            run = p.add_run(ctx_before)
            run.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        add_left_border(p)
        p.add_run(" ".join(text.split()))

        if ctx_after:
            p = doc.add_paragraph()
            run = p.add_run(ctx_after)
            run.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.save(path)


# ---------------------------------------------------------------------------
# Kindle clippings
# ---------------------------------------------------------------------------

def _parse_clippings(path):
    """Parse Kindle My Clippings.txt into a list of book dicts.

    Each dict: {"title": str, "author": str, "highlights": [{"text": str, "location": str|None}]}
    """
    try:
        with open(path, "r", encoding="utf-8-sig") as f:
            content = f.read()
    except Exception as e:
        raise Exception(f"Could not read clippings file: {e}")

    books = {}
    for entry in content.split("=========="):
        # Strip whitespace and BOM characters that appear mid-file on some firmware
        entry_stripped = entry.strip().replace("﻿", "")
        lines = [l.strip() for l in entry_stripped.splitlines() if l.strip()]
        if len(lines) < 2:
            continue

        title_line = lines[0]
        meta_line = lines[1]

        if "Your Highlight" not in meta_line:
            continue

        m = re.match(r"^(.*?)\s*\(([^)]*)\)\s*$", title_line)
        if m:
            title, author = m.group(1).strip(), m.group(2).strip()
        else:
            title, author = title_line.strip(), ""

        # Parse location/page from metadata line
        page_m = re.search(r"\bpage\s+(\d+)", meta_line, re.IGNORECASE)
        loc_m = re.search(r"\blocation\s+([\d\-]+)", meta_line, re.IGNORECASE)
        if page_m:
            location = f"Page {page_m.group(1)}"
            if loc_m:
                location += f" · Location {loc_m.group(1)}"
        elif loc_m:
            location = f"Location {loc_m.group(1)}"
        else:
            location = None

        # Text is everything after the metadata line
        text = " ".join(l for l in lines[2:] if l).strip()

        if not text:
            continue

        if title not in books:
            books[title] = {"title": title, "author": author, "highlights": []}
        books[title]["highlights"].append({"text": text, "location": location})

    return list(books.values())


def cmd_kindle_list(clippings_path):
    try:
        books = _parse_clippings(clippings_path)
    except Exception as e:
        _err(str(e))
        return

    result = []
    for b in books:
        if not b["highlights"]:
            continue
        result.append({
            "title": b["title"],
            "author": b["author"],
            "highlight_count": len(b["highlights"]),
        })
    result.sort(key=lambda x: x["title"].lower())
    _ok({"books": result})


def cmd_kindle_extract(clippings_path, book_title, output_path):
    try:
        books = _parse_clippings(clippings_path)
    except Exception as e:
        _err(str(e))
        return

    book = next((b for b in books if b["title"] == book_title), None)
    if book is None:
        _err(f"Book not found: {book_title}")
        return

    hits = [(h["location"], h["text"], "", "") for h in book["highlights"]]
    _write_docx(output_path, book_title, hits)
    _ok({"highlight_count": len(hits)})


# ---------------------------------------------------------------------------
# Output helpers
# ---------------------------------------------------------------------------

def _ok(extra=None):
    payload = {"status": "ok"}
    if extra:
        payload.update(extra)
    print(json.dumps(payload))


def _err(msg):
    print(json.dumps({"status": "error", "message": msg}))
    sys.exit(1)


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    if len(sys.argv) < 2:
        _err("Usage: extract.py <list|extract|kindle-list|kindle-extract> [args...]")

    command = sys.argv[1]

    try:
        if command == "list":
            if len(sys.argv) < 3:
                _err("list requires <host>")
            cmd_list(sys.argv[2])
        elif command == "extract":
            if len(sys.argv) < 5:
                _err("extract requires <host> <uuid> <output_path>")
            title = sys.argv[5] if len(sys.argv) > 5 else "Untitled"
            cmd_extract(sys.argv[2], sys.argv[3], sys.argv[4], title)
        elif command == "kindle-list":
            if len(sys.argv) < 3:
                _err("kindle-list requires <clippings_path>")
            cmd_kindle_list(sys.argv[2])
        elif command == "kindle-extract":
            if len(sys.argv) < 5:
                _err("kindle-extract requires <clippings_path> <book_title> <output_path>")
            cmd_kindle_extract(sys.argv[2], sys.argv[3], sys.argv[4])
        else:
            _err(f"Unknown command: {command}")
    except Exception as e:
        _err(str(e) + "\n" + traceback.format_exc())
